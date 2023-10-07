#!/usr/bin/env python3
"""
CV DOCX Generator - Abdelhak Bourdim
======================================

Generates DOCX CVs using python-docx with professional formatting.

Prerequisites:
    pip install python-docx

Usage:
    python3 generate_docx.py          # Generate both versions
    python3 generate_docx.py a        # Generate Version A only
    python3 generate_docx.py b        # Generate Version B only

Output:
    cv_abdelhak_bourdim_VERSION_A.docx  (Workshop-DIY in sidebar-style section)
    cv_abdelhak_bourdim_VERSION_B.docx  (Workshop-DIY as main experience)
"""

import sys
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

DIR = os.path.dirname(os.path.abspath(__file__))

# ── Colors ──
DARK_BLUE = RGBColor(0x1A, 0x3C, 0x6E)
MEDIUM_BLUE = RGBColor(0x2B, 0x57, 0x97)
BLACK = RGBColor(0x2D, 0x2D, 0x2D)
GRAY = RGBColor(0x55, 0x55, 0x55)
LIGHT_GRAY = RGBColor(0x8A, 0x9A, 0xAA)


def setup_doc():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.2)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(9.5)
    style.font.color.rgb = BLACK
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(2)
    style.paragraph_format.line_spacing = 1.1
    return doc


def add_heading_line(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10 if level == 1 else 6)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(12 if level == 1 else 10.5)
    run.font.color.rgb = DARK_BLUE
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="6" w:space="1" w:color="1A3C6E"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)
    return p


def add_company_header(doc, company, location, dates):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(f"{company} — {location}")
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = MEDIUM_BLUE
    p.add_run("\t")
    run_date = p.add_run(dates)
    run_date.font.size = Pt(9)
    run_date.font.color.rgb = GRAY
    run_date.italic = True
    tab_stops = p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Cm(17.4), alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    return p


def add_project_line(doc, project, intervention):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(f"Projet : {project}")
    run.bold = True
    run.font.size = Pt(9.5)
    run.font.color.rgb = BLACK
    run = p.add_run(f" — {intervention}")
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = GRAY
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.left_indent = Cm(1.2)
    p.clear()
    run = p.add_run(text)
    run.font.size = Pt(9)
    run.font.color.rgb = BLACK
    return p


def add_env_line(doc, env_text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run("Env. : ")
    run.bold = True
    run.font.size = Pt(8.5)
    run.font.color.rgb = LIGHT_GRAY
    run = p.add_run(env_text)
    run.font.size = Pt(8.5)
    run.font.color.rgb = LIGHT_GRAY
    return p


def remove_table_borders(table):
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = parse_xml(
                f'<w:tcBorders {nsdecls("w")}>'
                f'  <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                f'  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                f'  <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                f'  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                f'</w:tcBorders>'
            )
            tcPr.append(tcBorders)


def add_header(doc, include_workshop_mention=True):
    """Add name, title, contact, profile."""
    # Name
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("ABDELHAK BOURDIM")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = DARK_BLUE

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("Ingénieur Logiciel Embarqué Senior")
    run.font.size = Pt(13)
    run.font.color.rgb = MEDIUM_BLUE
    run = p.add_run("  —  C/C++ | RTOS | Linux Embarqué")
    run.font.size = Pt(11)
    run.font.color.rgb = GRAY

    # Contact
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run("abdelhak.bourdim@gmail.com  |  +33 6 19 51 51 73  |  Région parisienne")
    run.font.size = Pt(9)
    run.font.color.rgb = GRAY

    # LinkedIn + site
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("linkedin.com/in/abdelhak-bourdim-96416122  |  workshop-diy.org")
    run.font.size = Pt(8.5)
    run.font.color.rgb = MEDIUM_BLUE

    # Profile
    add_heading_line(doc, "Profil")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)

    run = p.add_run(
        "Ingénieur embarqué avec plus de 20 ans d'expérience dans le développement de logiciels "
        "bas niveau et applicatifs pour des systèmes critiques. Expertise en développement de drivers, "
        "protocoles de communication (CAN, BLE, USB, TCP/IP) et RTOS sur architectures ARM Cortex, "
        "STM32, Microchip et NXP. Interventions dans les secteurs aéronautique (DO-178, Safran, Thales), "
        "médical (IEC 62304, GE Healthcare, PK Vitality), automobile (Valeo, Arkamys) et IoT. "
        "Développement conforme MISRA C."
    )
    run.font.size = Pt(9.5)
    run.font.color.rgb = BLACK

    if include_workshop_mention:
        run = p.add_run(
            " Fondateur de Workshop-DIY, association dédiée à la robotique éducative et au prototypage électronique."
        )
        run.font.size = Pt(9.5)
        run.font.color.rgb = BLACK


def add_skills_table(doc, include_esp32=False):
    """Add skills table."""
    add_heading_line(doc, "Compétences techniques")

    mcu_text = (
        "ARM Cortex (M0, M4, A8), STM32L4, NORDIC NRF52, ESP32, Microchip (PIC, dsPIC), "
        "NXP LPC, TI DSP320/Concerto, Infineon XE167, Hitachi H8S, NEC V850, Arduino"
    ) if include_esp32 else (
        "ARM Cortex (M0, M4, A8), STM32L4, NORDIC NRF52, ESP32, Microchip (PIC, dsPIC), "
        "NXP LPC, TI DSP320/Concerto, Infineon XE167, Hitachi H8S, Arduino"
    )

    gestion_text = "Git, GitHub, Synergy, Jira, Polarion, Doxygen" if include_esp32 else "Git, Synergy, Jira, Polarion, Doxygen"

    skills_data = [
        ("Langages", "C, C++, Python, Bash, Assembleur, Perl, Awk"),
        ("OS / RTOS", "Linux embarqué, VxWorks, FreeRTOS, UCOSII, Tnkernel"),
        ("Microcontrôleurs", mcu_text),
        ("Protocoles / Bus", "CAN, CANOpen, SPI, I2C, RS232, RS485, UART, USB, BLE, AFDX, Ethernet, TCP/UDP, PWM, ADC, DAC"),
        ("Stacks", "TCP/IP, USB (Host/Device/HID/CCID), Bluetooth Low Energy, CANOpen, RFID, GStreamer, IRDA"),
        ("Outils embarqué", "Lauterbach TRACE32, IAR, Keil, MPLAB, CCS Studio, Green Hills, Yocto, Buildroot, CubeMx, TouchGFX"),
        ("Gestion / CI", gestion_text),
        ("Normes", "DO-178, MISRA C, IEC 62304, Cycle en V, Agile / Scrum"),
        ("Langues", "Français (courant), Anglais (professionnel), Arabe (courant)"),
    ]

    table = doc.add_table(rows=len(skills_data), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = True

    for i, (label, value) in enumerate(skills_data):
        cell_label = table.cell(i, 0)
        cell_value = table.cell(i, 1)
        p = cell_label.paragraphs[0]
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run(label)
        run.bold = True
        run.font.size = Pt(8.5)
        run.font.color.rgb = DARK_BLUE
        p = cell_value.paragraphs[0]
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run(value)
        run.font.size = Pt(8.5)
        run.font.color.rgb = BLACK
        cell_label.width = Cm(3.5)

    remove_table_borders(table)


def add_common_experiences(doc):
    """Shared experience blocks (Safran through Arkamys)."""

    # SAFRAN
    add_company_header(doc, "Safran", "Créteil", "Juillet 2024 – Mai 2025")
    add_project_line(doc, "SSPC", "Mise à jour des drivers bas niveau")
    add_bullet(doc, "Développement et mise à jour des drivers bas niveau en C pour le système SSPC (Solid State Power Controller)")
    add_bullet(doc, "Mise en place d'outils de logging pour le diagnostic et la traçabilité des échanges")
    add_bullet(doc, "Rédaction et exécution des tests de validation sur cible avec Lauterbach TRACE32")
    add_bullet(doc, "Scripting Bash/Awk pour l'automatisation des processus de test")
    add_env_line(doc, "C, Lauterbach TRACE32, Bash, Awk")

    # ASTERIOS
    add_company_header(doc, "Asterios Technologies", "Massy", "Août 2023 – Février 2024")
    add_project_line(doc, "T1042SOM", "Rédaction des procédures de tests logiciels")
    add_bullet(doc, "Développement en C des tests de validation pour le module SOM basé sur le processeur T1042")
    add_bullet(doc, "Analyse des errata du processeur et identification des contournements logiciels")
    add_bullet(doc, "Rédaction de la documentation de test conforme aux exigences du projet")
    add_env_line(doc, "C, Bash, Python, Lauterbach TRACE32, Polarion")

    # THALES GENNEVILLIERS
    add_company_header(doc, "Thales", "Gennevilliers", "Mai 2022 – Juillet 2023")
    add_project_line(doc, "LPM + N-MMR", "Logiciel de préparation de mission & New Multi Mode Receiver")
    add_bullet(doc, "Développement des couches applicatives sous Linux embarqué en C/C++ pour 2 projets radio défense")
    add_bullet(doc, "Scripts de gestion de configuration et automatisation avec Bash, Curl et Tshark")
    add_bullet(doc, "Développement et exécution des tests HLT (High Level Tests)")
    add_bullet(doc, "Correction de bugs et analyse des protocoles réseau (Websockets, Tshark)")
    add_env_line(doc, "Linux embarqué, C/C++, Bash, Websockets, Curl, Tshark, Python")

    # VALEO
    add_company_header(doc, "Valeo", "Créteil", "Mars – Avril 2022")
    add_project_line(doc, "Audit des tests", "Audit pour Valeo Egypt")
    add_bullet(doc, "Audit des processus de tests (NI Mastre, Vector CANoe), propositions d'automatisation")
    add_bullet(doc, "Revue des exigences client et système, rédaction du rapport d'audit")
    add_env_line(doc, "NI Mastre, Vector CANoe")

    # ARKAMYS
    add_company_header(doc, "Arkamys", "Levallois-Perret", "Juillet 2021 – Février 2022")
    add_project_line(doc, "GStreamer Framework", "Plugins audio, tests unitaires & d'intégration")
    add_bullet(doc, "Développement de plugins GStreamer et drivers CAN (ELM327)")
    add_bullet(doc, "Tests unitaires & d'intégration avec Valgrind, build Yocto/Buildroot")
    add_env_line(doc, "Linux embarqué, C/C++, Yocto, Buildroot, GStreamer, Valgrind, CAN, Python, Jira")


def add_remaining_experiences(doc):
    """Thales Velizy through GE Healthcare."""

    # THALES VELIZY
    add_company_header(doc, "Thales", "Vélizy-Villacoublay", "Novembre 2020 – Juin 2021")
    add_project_line(doc, "ATO — Autonomous Train Operation", "IPC communication & socket programming")
    add_bullet(doc, "Migration de l'application de 32 bits vers 64 bits sur plateforme Linux embarqué")
    add_bullet(doc, "Définition et implémentation des protocoles de communication IPC")
    add_bullet(doc, "Développement d'outils de diagnostique réseau et tests d'intégration")
    add_bullet(doc, "Analyse des trames avec Wireshark, Tshark, Pscapy et Pyshark")
    add_env_line(doc, "Linux embarqué, C, UDP, TCP, Wireshark, Tshark, Pscapy, Python, Jira")

    # PK VITALITY
    add_company_header(doc, "PK Vitality", "Paris", "Juin – Septembre 2020")
    add_project_line(doc, "K'Watch CGM — Continuous Glucose Monitor", "Drivers BLE, FreeRTOS, STM32 TouchGFX")
    add_bullet(doc, "Développement du HAL pour l'AFE (AD5940) et du PMIC (DA9070)")
    add_bullet(doc, "Driver QSPI, gestion des logs patient, BLE avec NORDIC NRF52840")
    add_bullet(doc, "Écran tactile STM32 TouchGFX, gestion tâches FreeRTOS, mode low power")
    add_bullet(doc, "Tests in vitro du capteur et documentation Doxygen (conforme IEC 62304)")
    add_env_line(doc, "STM32L4S9, NORDIC NRF52840, C, C++, FreeRTOS, BLE, TouchGFX, Doxygen")

    # ZODIAC
    add_company_header(doc, "Zodiac Aerospace (Safran)", "Montreuil", "Février 2017 – Mars 2020")
    add_project_line(doc, "GENOME / CORAC / SSPC A350 AC9", "Drivers bas niveau, diagnostic, couches applicatives")
    add_bullet(doc, "Adaptation des couches bas niveau : drivers CAN, RS485, mémoire thermique (dsPIC33)")
    add_bullet(doc, "Implémentation de la passerelle uAFDX vers CAN et fonctions de diagnostique")
    add_bullet(doc, "Développement et optimisation des couches applicatives SSPC A350")
    add_bullet(doc, "Tests de validation et documentation conforme aux standards aéronautiques")
    add_env_line(doc, "Microchip dsPIC33, MPLAB, TI Concerto, CCS Studio, C")

    # SORIN
    add_company_header(doc, "Sorin", "Clamart", "Novembre 2014 – Avril 2015")
    add_project_line(doc, "Platinium — Module BLE pour Pacemaker", "Étude BLE pour implant cardiaque")
    add_bullet(doc, "Validation du module Bluetooth Smart, développement du Serial Port Service Profile")
    add_bullet(doc, "Protocole de communication smartphone et application Android")
    add_bullet(doc, "Interface hardware et préparation des échantillons pour implantation")
    add_env_line(doc, "Cortex M0, Keil, C, Android (Eclipse)")

    # GE HEALTHCARE
    add_company_header(doc, "General Electric Healthcare", "Buc", "Novembre 2012 – Juillet 2014")
    add_project_line(doc, "ZYX — Générateur Rayons X", "Intégration CANOpen sur cartes RCB et CORTYX")
    add_bullet(doc, "Intégration pile CANOpen (slave/master), développement drivers et service CAN")
    add_bullet(doc, "Portage du noyau CMX-RTX, gestion fichiers, scripts TCL de test")
    add_bullet(doc, "Support technique équipe USA et déplacements sur site aux États-Unis")
    add_env_line(doc, "TI TMS320, XE167, ARM Cortex A8, VxWorks, C, C++")


def add_older_experiences(doc):
    """Condensed older experiences."""
    add_heading_line(doc, "Expériences antérieures")

    older = [
        ("Enedis — Nanterre", "Mai 2015 – Sep. 2016",
         "Compteurs Linky (smart grids) — Diagnostique des concentrateurs, automatisation (Arduino, Raspberry Pi, Python, Bash)"),
        ("Gimelec — Algérie", "Avr. – Août 2012",
         "Generator Control Unit — Produit fini, développement complet HW+SW (PIC16F628, RS485)"),
        ("Photomaton / BCM / Vertical M2M / PGS — Paris", "2010 – 2011",
         "Systèmes de paiement et télé relève — Développement embarqué + PC sur NXP LPC (USB, RFID, GSM/GPRS, Bluetooth)"),
        ("SFR Team — Paris", "Sep. 2009 – Fév. 2010",
         "Télé relève autonome pour citernes LPG — GSM/GPRS, Microchip PIC24F"),
        ("Cartadis — Fontenay-sous-Bois", "Oct. 2001 – Juin 2009",
         "8 ans — Terminaux de contrôle d'accès et paiement : drivers, stacks USB/TCP/IP, cryptage DES/RC4 (NEC V853, Hitachi H8S, NXP LPC, ATMEL AT91)"),
    ]

    for company, dates, desc in older:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run(company)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = MEDIUM_BLUE
        run = p.add_run(f"  ({dates})")
        run.font.size = Pt(8.5)
        run.font.color.rgb = GRAY
        run.italic = True
        run = p.add_run(f"\n{desc}")
        run.font.size = Pt(8.5)
        run.font.color.rgb = BLACK


def add_education(doc):
    """Add education section."""
    add_heading_line(doc, "Formation")
    formations = [
        ("2001", "Ingénieur de conception de systèmes électroniques et informatiques — AFPA, Champs-sur-Marne"),
        ("1997", "Diplôme d'instrumentation biomédicale (3ème cycle) — Hôpital Tenon, Paris"),
        ("1994", "Diplôme d'ingénieur en électronique, option communication — Boumerdès, Algérie"),
    ]
    for year, desc in formations:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run(f"{year}  ")
        run.bold = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = DARK_BLUE
        run = p.add_run(desc)
        run.font.size = Pt(9)
        run.font.color.rgb = BLACK


# ═══════════════════════════════════════════
# VERSION A — Workshop-DIY as sidebar mention
# ═══════════════════════════════════════════
def generate_version_a():
    doc = setup_doc()
    add_header(doc, include_workshop_mention=True)
    add_skills_table(doc, include_esp32=True)
    add_heading_line(doc, "Expérience professionnelle")
    add_common_experiences(doc)
    add_remaining_experiences(doc)
    add_older_experiences(doc)
    add_education(doc)

    out = os.path.join(DIR, "cv_abdelhak_bourdim_VERSION_A.docx")
    doc.save(out)
    return out


# ═══════════════════════════════════════════
# VERSION B — Workshop-DIY as main experience
# ═══════════════════════════════════════════
def generate_version_b():
    doc = setup_doc()
    add_header(doc, include_workshop_mention=True)
    add_skills_table(doc, include_esp32=True)
    add_heading_line(doc, "Expérience professionnelle")

    # WORKSHOP-DIY first
    add_company_header(doc, "Workshop-DIY (Association)", "Chelles (77)", "Depuis 2020 — En cours")
    add_project_line(doc, "Fondateur — Robotique & STEM",
                     "Conception de kits robotiques, ateliers pédagogiques, projets IA & prototypage")
    add_bullet(doc, "Création et gestion d'une association proposant 18+ ateliers robotique & STEM pour jeunes (6-12+ ans)")
    add_bullet(doc, "Conception d'un kit robot ESP32 complet (WiFi, BLE, servos, capteurs) commercialisé en ligne")
    add_bullet(doc, "Projets IA : détection faciale, hand tracking, synthèse vocale (Python, OpenCV)")
    add_bullet(doc, "Ateliers Arduino, micro:bit, gravure laser, impression 3D — site multilingue workshop-diy.org")
    add_env_line(doc, "ESP32, Arduino, micro:bit, Python, OpenCV, C/C++, HTML/CSS/JS, GitHub")

    add_common_experiences(doc)
    add_remaining_experiences(doc)
    add_older_experiences(doc)
    add_education(doc)

    out = os.path.join(DIR, "cv_abdelhak_bourdim_VERSION_B.docx")
    doc.save(out)
    return out


# ═══════════════════════════════════════════
if __name__ == "__main__":
    choice = sys.argv[1].lower() if len(sys.argv) > 1 else "all"

    print("CV DOCX Generator")
    print("=" * 40)

    if choice in ("a", "all"):
        out = generate_version_a()
        size_kb = os.path.getsize(out) / 1024
        print(f"  [OK] Version A -> {out} ({size_kb:.1f} KB)")

    if choice in ("b", "all"):
        out = generate_version_b()
        size_kb = os.path.getsize(out) / 1024
        print(f"  [OK] Version B -> {out} ({size_kb:.1f} KB)")

    print("\nDone!")
