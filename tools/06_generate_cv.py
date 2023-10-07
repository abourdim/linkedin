#!/usr/bin/env python3
"""Generate improved CV for Abdelhak Bourdim as DOCX."""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

doc = Document()

# ── Page margins ──
for section in doc.sections:
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.2)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

# ── Style helpers ──
DARK_BLUE = RGBColor(0x1A, 0x3C, 0x6E)
MEDIUM_BLUE = RGBColor(0x2B, 0x57, 0x97)
BLACK = RGBColor(0x2D, 0x2D, 0x2D)
GRAY = RGBColor(0x55, 0x55, 0x55)

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(9.5)
style.font.color.rgb = BLACK
style.paragraph_format.space_before = Pt(0)
style.paragraph_format.space_after = Pt(2)
style.paragraph_format.line_spacing = 1.1


def add_heading_line(text, level=1):
    """Add a heading with a bottom border line."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10 if level == 1 else 6)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(12 if level == 1 else 10.5)
    run.font.color.rgb = DARK_BLUE
    # Bottom border
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="6" w:space="1" w:color="1A3C6E"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)
    return p


def add_company_header(company, location, dates):
    """Add company name with dates right-aligned."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after = Pt(1)
    # Company + location
    run = p.add_run(f"{company} — {location}")
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = MEDIUM_BLUE
    # Tab stop for right-aligned dates
    p.add_run("\t")
    run_date = p.add_run(dates)
    run_date.font.size = Pt(9)
    run_date.font.color.rgb = GRAY
    run_date.italic = True
    # Right tab stop
    tab_stops = p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Cm(17.4), alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    return p


def add_project_line(project, intervention):
    """Add project name and intervention description."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(f"Projet : {project}")
    run.bold = True
    run.font.size = Pt(9.5)
    run.font.color.rgb = BLACK
    run = p.add_run(f" — {intervention}")
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = GRAY
    return p


def add_bullet(text):
    """Add a bullet point."""
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.left_indent = Cm(1.2)
    # Clear default run and add formatted one
    p.clear()
    run = p.add_run(text)
    run.font.size = Pt(9)
    run.font.color.rgb = BLACK
    return p


def add_env_line(env_text):
    """Add environment/tech line."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run("Env. : ")
    run.bold = True
    run.font.size = Pt(8.5)
    run.font.color.rgb = GRAY
    run = p.add_run(env_text)
    run.font.size = Pt(8.5)
    run.font.color.rgb = GRAY
    return p


# =====================================================================
# PAGE 1 — HEADER
# =====================================================================

# Name
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(2)
run = p.add_run("ABDELHAK BOURDIM")
run.bold = True
run.font.size = Pt(20)
run.font.color.rgb = DARK_BLUE

# Title
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after = Pt(4)
run = p.add_run("Ingénieur Logiciel Embarqué Senior")
run.font.size = Pt(13)
run.font.color.rgb = MEDIUM_BLUE
run = p.add_run("  —  C/C++ | RTOS | Linux Embarqué")
run.font.size = Pt(11)
run.font.color.rgb = GRAY

# Contact line
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(2)
run = p.add_run("abdelhak.bourdim@gmail.com  |  +33 6 19 51 51 73  |  Région parisienne")
run.font.size = Pt(9)
run.font.color.rgb = GRAY

# ── PROFIL ──
add_heading_line("Profil")
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(2)
run = p.add_run(
    "Ingénieur embarqué avec plus de 20 ans d'expérience dans le développement de logiciels "
    "bas niveau et applicatifs pour des systèmes critiques. Expertise en développement de drivers, "
    "protocoles de communication (CAN, BLE, USB, TCP/IP) et RTOS sur architectures ARM Cortex, "
    "STM32, Microchip et NXP. Interventions dans les secteurs aéronautique (DO-178, Safran, Thales), "
    "médical (IEC 62304, GE Healthcare, PK Vitality), automobile (Valeo, Arkamys) et IoT. "
    "Développement conforme MISRA C."
)
run.font.size = Pt(9.5)
run.font.color.rgb = BLACK

# ── COMPÉTENCES TECHNIQUES ──
add_heading_line("Compétences techniques")

skills_data = [
    ("Langages", "C, C++, Python, Bash, Assembleur, Perl, Awk"),
    ("OS / RTOS", "Linux embarqué, VxWorks, FreeRTOS, UCOSII, Tnkernel"),
    ("Microcontrôleurs",
     "ARM Cortex (M0, M4, A8), STM32L4, NORDIC NRF52, Microchip (PIC, dsPIC), "
     "NXP LPC, TI DSP320/Concerto, Infineon XE167, Hitachi H8S, NEC V850"),
    ("Protocoles / Bus", "CAN, CANOpen, SPI, I2C, RS232, RS485, UART, USB, BLE, AFDX, Ethernet, TCP/UDP, PWM, ADC, DAC"),
    ("Stacks", "TCP/IP, USB (Host/Device/HID/CCID), Bluetooth Low Energy, CANOpen, RFID, GStreamer, IRDA"),
    ("Outils embarqué", "Lauterbach TRACE32, IAR, Keil, MPLAB, CCS Studio, Green Hills, Yocto, Buildroot, CubeMx, TouchGFX"),
    ("Gestion / CI", "Git, Synergy, Jira, Polarion, Doxygen"),
    ("Normes", "DO-178, MISRA C, IEC 62304, Cycle en V, Agile / Scrum"),
    ("Langues", "Français (courant), Anglais (professionnel), Arabe (courant)"),
]

table = doc.add_table(rows=len(skills_data), cols=2)
table.alignment = WD_TABLE_ALIGNMENT.LEFT
table.autofit = True

for i, (label, value) in enumerate(skills_data):
    cell_label = table.cell(i, 0)
    cell_value = table.cell(i, 1)

    # Label cell
    p = cell_label.paragraphs[0]
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(label)
    run.bold = True
    run.font.size = Pt(8.5)
    run.font.color.rgb = DARK_BLUE

    # Value cell
    p = cell_value.paragraphs[0]
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(value)
    run.font.size = Pt(8.5)
    run.font.color.rgb = BLACK

    # Set label column width
    cell_label.width = Cm(3.5)

# Remove table borders
for row in table.rows:
    for cell in row.cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = parse_xml(
            f'<w:tcBorders {nsdecls("w")}>'
            f'  <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            f'  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            f'  <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            f'  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            f'</w:tcBorders>'
        )
        tcPr.append(tcBorders)

# =====================================================================
# EXPÉRIENCE PROFESSIONNELLE
# =====================================================================
add_heading_line("Expérience professionnelle")

# ── SAFRAN ──
add_company_header("Safran", "Créteil", "Juillet 2024 – Mai 2025")
add_project_line("SSPC", "Mise à jour des drivers bas niveau")
add_bullet("Développement et mise à jour des drivers bas niveau en C pour le système SSPC (Solid State Power Controller)")
add_bullet("Mise en place d'outils de logging pour le diagnostic et la traçabilité des échanges")
add_bullet("Rédaction et exécution des tests de validation sur cible avec Lauterbach TRACE32")
add_bullet("Scripting Bash/Awk pour l'automatisation des processus de test")
add_env_line("C, Lauterbach TRACE32, Bash, Awk")

# ── ASTERIOS ──
add_company_header("Asterios Technologies", "Massy", "Août 2023 – Février 2024")
add_project_line("T1042SOM", "Rédaction des procédures de tests logiciels")
add_bullet("Développement en C des tests de validation pour le module SOM basé sur le processeur T1042")
add_bullet("Analyse des errata du processeur et identification des contournements logiciels")
add_bullet("Rédaction de la documentation de test conforme aux exigences du projet")
add_env_line("C, Bash, Python, Lauterbach TRACE32, Polarion")

# ── THALES GENNEVILLIERS ──
add_company_header("Thales", "Gennevilliers", "Mai 2022 – Juillet 2023")
add_project_line("LPM + N-MMR",
                 "Logiciel de préparation de mission & New Multi Mode Receiver")
add_bullet("Développement des couches applicatives sous Linux embarqué en C/C++ pour 2 projets radio défense")
add_bullet("Mise à jour de l'application conformément aux évolutions des spécifications")
add_bullet("Scripts de gestion de configuration et automatisation avec Bash, Curl et Tshark")
add_bullet("Développement et exécution des tests HLT (High Level Tests)")
add_bullet("Correction de bugs et analyse des protocoles réseau (Websockets, Tshark)")
add_env_line("Linux embarqué, C/C++, Bash, Websockets, Curl, Tshark, Python")

# ── VALEO ──
add_company_header("Valeo", "Créteil", "Mars 2022 – Avril 2022")
add_project_line("Audit des tests", "Audit pour Valeo Egypt")
add_bullet("Audit des processus de tests unitaires, d'intégration et des outils utilisés (NI Mastre, CANoe)")
add_bullet("Propositions d'automatisation des tests manuels et d'amélioration des temps d'exécution")
add_bullet("Revue des exigences client et système, rédaction du rapport d'audit")
add_env_line("NI Mastre, Vector CANoe")

# ── ARKAMYS ──
add_company_header("Arkamys", "Levallois-Perret", "Juillet 2021 – Février 2022")
add_project_line("GStreamer Framework", "Plugins audio, tests unitaires & d'intégration")
add_bullet("Développement de plugins GStreamer pour le traitement audio embarqué")
add_bullet("Développement des drivers CAN (ELM327) et des outils de diagnostique")
add_bullet("Rédaction des plans de validation et exécution des tests unitaires & d'intégration")
add_bullet("Build system avec Yocto/Buildroot, analyse mémoire avec Valgrind")
add_env_line("Linux embarqué, C/C++, Yocto, Buildroot, GStreamer, Valgrind, CAN ELM327, Python, Bash, Jira")

# ── THALES VELIZY ──
add_company_header("Thales", "Vélizy-Villacoublay", "Novembre 2020 – Juin 2021")
add_project_line("ATO — Autonomous Train Operation", "IPC communication & socket programming")
add_bullet("Migration de l'application de 32 bits vers 64 bits sur plateforme Linux embarqué")
add_bullet("Définition et implémentation des protocoles de communication inter-processus (IPC)")
add_bullet("Développement d'outils de diagnostique réseau et tests d'intégration")
add_bullet("Analyse des trames avec Wireshark, Tshark, Pscapy et Pyshark")
add_env_line("Linux embarqué, C, UDP, TCP, Wireshark, Tshark, Pscapy, Pyshark, Python, Awk, Jira")

# ── PK VITALITY ──
add_company_header("PK Vitality", "Paris", "Juin 2020 – Septembre 2020")
add_project_line("K'Watch CGM — Continuous Glucose Monitor",
                 "Drivers BLE, FreeRTOS, STM32 TouchGFX")
add_bullet("Développement du HAL pour l'AFE (Analog Front End AD5940) et du PMIC (DA9070)")
add_bullet("Driver QSPI et gestion des logs patient, communication BLE avec NORDIC NRF52840")
add_bullet("Gestion de l'écran tactile avec STM32 TouchGFX, CubeMx et Touch Designer")
add_bullet("Gestion des tâches FreeRTOS et optimisation du mode low power")
add_bullet("Tests in vitro du capteur et documentation Doxygen (conforme IEC 62304)")
add_env_line("STM32L4S9, NORDIC NRF52840, C, C++, FreeRTOS, BLE, TouchGFX, Doxygen")

# ── ZODIAC AEROSPACE (SAFRAN) ──
add_company_header("Zodiac Aerospace (Safran)", "Montreuil", "Février 2017 – Mars 2020")
add_project_line("GENOME / CORAC / SSPC A350 AC9",
                 "Drivers bas niveau, diagnostic, couches applicatives")
add_bullet("Adaptation des couches bas niveau et implémentation des drivers CAN, RS485 et mémoire thermique (Microchip dsPIC33)")
add_bullet("Implémentation de la couche passerelle uAFDX vers CAN et des fonctions de diagnostique")
add_bullet("Développement et optimisation des couches applicatives pour le système SSPC A350")
add_bullet("Réalisation des tests de validation et documentation conforme aux standards internes")
add_env_line("Microchip dsPIC33, MPLAB, TI Concerto, CCS Studio, C")

# =====================================================================
# PAGE 3 — EXPÉRIENCES ANTÉRIEURES + FORMATION
# =====================================================================
add_heading_line("Expériences antérieures")

older_experiences = [
    ("Sorin — Clamart", "Nov. 2014 – Avr. 2015",
     "Étude BLE pour pacemaker (Cortex M0, Keil, Android)"),
    ("General Electric Healthcare — Buc", "Nov. 2012 – Juil. 2014",
     "Intégration CANOpen, drivers CAN, portage RTOS (VxWorks, CMX-RTX, ARM Cortex A8, TI TMS320)"),
    ("Gimelec — Algérie", "Avr. 2012 – Août 2012",
     "Produit fini : Generator Control Unit (PIC16F628, RS485)"),
    ("Photomaton — Paris", "Mai 2011 – Oct. 2011",
     "Système de paiement par monnayeur (NXP LPC1754, USB, IAR, SolidWorks)"),
    ("BCM — Paris", "Jan. 2011 – Avr. 2011",
     "Système de paiement par cartes Mifare RFID (NXP LPC1754, ISO/IEC 14443-3)"),
    ("Vertical M2M — Paris", "Juin 2010 – Nov. 2010",
     "Télé relève GSM/GPRS de compteurs (NXP LPC1763, RF 433MHz)"),
    ("PGS — Lyon", "Mars 2010 – Mai 2010",
     "Gestion de parc de photocopieurs via Bluetooth (NXP LPC1343)"),
    ("SFR Team — Paris", "Sep. 2009 – Fév. 2010",
     "Télé relève autonome pour citernes LPG (Microchip PIC24F, GSM/GPRS)"),
    ("Cartadis — Fontenay-sous-Bois", "Oct. 2001 – Juin 2009",
     "8 ans — Terminaux de contrôle d'accès et de paiement : drivers, applicatifs embarqués, "
     "stacks USB/TCP/IP, cryptage (NEC V853, Hitachi H8S, NXP LPC, ATMEL AT91)"),
    ("Datus — Boulogne-Billancourt", "Juin 2001 – Août 2001",
     "Stage — Pilotage de moteur par PC, maquette TGV SNCF (Microchip PIC16F873)"),
]

for company, dates, desc in older_experiences:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(1)

    run = p.add_run(f"{company}")
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = MEDIUM_BLUE

    run = p.add_run(f"  ({dates})")
    run.font.size = Pt(8.5)
    run.font.color.rgb = GRAY
    run.italic = True

    run = p.add_run(f"\n{desc}")
    run.font.size = Pt(8.5)
    run.font.color.rgb = BLACK

# ── FORMATION ──
add_heading_line("Formation")

formations = [
    ("2001", "Ingénieur de conception de systèmes électroniques et informatiques — AFPA, Champs-sur-Marne"),
    ("1997", "Diplôme d'instrumentation biomédicale (3ème cycle) — Hôpital Tenon, Paris"),
    ("1994", "Diplôme d'ingénieur en électronique, option communication et télécommunication — Boumerdès, Algérie"),
]

for year, desc in formations:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(f"{year}  ")
    run.bold = True
    run.font.size = Pt(9.5)
    run.font.color.rgb = DARK_BLUE
    run = p.add_run(desc)
    run.font.size = Pt(9)
    run.font.color.rgb = BLACK

# ── Save ──
output_path = "/home/abdelhak/Downloads/02_cv/cv_abdelhak_bourdim_fr_2025_v2.docx"
doc.save(output_path)
print(f"CV saved to: {output_path}")
